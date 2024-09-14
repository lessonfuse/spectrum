from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from io import BytesIO
from PIL import Image

class ICPDocumentGenerator:
    def __init__(self, student):
        self.student = student
        self.document = Document()

    def generate(self):
        self._add_logo_and_header()
        self._add_student_info_table()
        self._add_learning_profile()
        self._add_developmental_areas()
        self._add_skills_strengths()
        self._add_accessible_learning_support()
        self._add_goals()
        self._add_intervention_services()
        self._add_supplementary_services()
        return self.document

    def _add_logo_and_header(self):
        if self.student.school_logo:
            logo = Image.open(BytesIO(self.student.school_logo.read()))
            logo_width = Cm(2.5)  # Adjust as needed
            logo_height = int(logo_width * (logo.height / logo.width))
            logo_paragraph = self.document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(self.student.school_logo, width=logo_width, height=Cm(logo_height))

        header = self.document.add_heading('INDIVIDUALIZED CURRICULUM PLAN', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_student_info_table(self):
        table = self.document.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'Name'
        cells[1].text = self.student.name

        cells = table.rows[0].cells
        cells[0].text = 'ID Card Number'
        cells[1].text = self.student.id_card_number

        cells = table.rows[1].cells
        cells[0].text = 'Age'
        cells[1].text = str(self.student.age)

        cells = table.rows[1].cells
        cells[0].text = 'IE Program'
        cells[1].text = self.student.ie_program

        cells = table.rows[2].cells
        cells[0].text = 'Index'
        cells[1].text = self.student.index

        cells = table.rows[2].cells
        cells[0].text = 'Date'
        cells[1].text = str(self.student.date_of_document)

    def _add_learning_profile(self):
        self.document.add_heading('Learning Profile', level=1)
        learning_profile = self.student.learning_profile
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        for field in learning_profile._meta.get_fields():
            if field.name not in ['id', 'student']:
                row = table.add_row()
                row.cells[0].text = field.verbose_name
                value = getattr(learning_profile, field.name)
                row.cells[1].text = dict(learning_profile.CONDITION_STATUS)[value] if isinstance(value, str) else str(value)

    def _add_learning_profile(self):
        self.document.add_heading('Learning Profile', level=1)
        learning_profile = self.student.learning_profile
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        for field in learning_profile._meta.get_fields():
            if field.name not in ['id', 'student']:
                row = table.add_row()
                row.cells[0].text = field.verbose_name
                value = getattr(learning_profile, field.name)
                row.cells[1].text = dict(learning_profile.CONDITION_STATUS)[value] if isinstance(value, str) else str(value)

    def _add_developmental_areas(self):
        self.document.add_heading('Developmental Areas', level=1)
        dev_areas = self.student.developmental_area
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        for field in dev_areas._meta.get_fields():
            if field.name not in ['id', 'student']:
                row = table.add_row()
                row.cells[0].text = field.verbose_name
                value = getattr(dev_areas, field.name)
                row.cells[1].text = 'Yes' if value else 'No' if isinstance(value, bool) else str(value)

    def _add_skills_strengths(self):
        self.document.add_heading('Skills & Strengths', level=1)
        skills = self.student.skills_strengths
        for field in skills._meta.get_fields():
            if field.name not in ['id', 'student']:
                self.document.add_paragraph(f'{field.verbose_name}: {getattr(skills, field.name)}')

    def _add_accessible_learning_support(self):
        self.document.add_heading('Accessible Learning Support', level=1)
        support = self.student.accessible_learning_support
        for field in support._meta.get_fields():
            if field.name not in ['id', 'student']:
                self.document.add_paragraph(f'{field.verbose_name}: {getattr(support, field.name)}')

    def _add_goals(self):
        self.document.add_heading('Measurable Goals', level=1)
        goals = self.student.goals.all()
        for goal in goals:
            table = self.document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            for field in goal._meta.get_fields():
                if field.name not in ['id', 'student']:
                    row = table.add_row()
                    row.cells[0].text = field.verbose_name
                    row.cells[1].text = str(getattr(goal, field.name))
            self.document.add_paragraph()  # Add space between goals

    def _add_intervention_services(self):
        self.document.add_heading('Intervention Services', level=1)
        services = self.student.intervention_services.all()
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'Program'
        headers[1].text = 'Frequency'
        headers[2].text = 'Duration'
        headers[3].text = 'Location'
        headers[4].text = 'Initiation Date'
        for service in services:
            row = table.add_row()
            row.cells[0].text = service.program
            row.cells[1].text = service.frequency
            row.cells[2].text = service.duration
            row.cells[3].text = service.location
            row.cells[4].text = str(service.initiation_date)

    def _add_supplementary_services(self):
        self.document.add_heading('Supplementary Services', level=1)
        services = self.student.supplementary_services.all()
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'Type of Support'
        headers[1].text = 'Time'
        headers[2].text = 'Duration'
        for service in services:
            row = table.add_row()
            row.cells[0].text = service.type_of_support
            row.cells[1].text = service.time
            row.cells[2].text = service.duration
