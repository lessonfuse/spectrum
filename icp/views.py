from django.forms import modelform_factory
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.urls import reverse_lazy
from django_flatpickr.schemas import FlatpickrOptions
from django_flatpickr.widgets import DatePickerInput
from docx import Document

from onlydjango.helpers.cbv import ODCreateView, ODListView, ODUpdateView, ODDeleteView, ODDetailView
from .models import (
    Student, Goal, InterventionService,
    SupplementaryService, GeneralInformation, LearningProfile,
    DevelopmentalArea, SkillsStrengths, AccessibleLearningSupport
)
from .forms import LearningProfileForm


def home(request):
    return render(request, 'icp/home.html')

class ICPListView(ODListView):
    model = Student
    context_object_name = 'students'
    template_name = 'icp/list.html'

class StudentDetailView(ODDetailView):
    model = Student
    template_name = 'icp/student_detail.html'

class StudentCreateView(ODCreateView):
    model = Student
    fields = ['name', 'id_card_number', 'ie_program', 'date_of_document']
    template_name = 'generic/create.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Create Student'
        return context

    def get_form_class(self):
        return modelform_factory(
            self.model,
            fields=self.fields,
            widgets={
                "date_of_document": DatePickerInput(
                    options=FlatpickrOptions(minDate="today"),
                ),
            },
        )

    def form_valid(self, form):
        self.object = form.save()
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('general_information', kwargs={'student_id': self.object.id})

    def post(self, request, *args, **kwargs):
        self.object = None
        return super().post(request, *args, **kwargs)

class StudentUpdateView(ODUpdateView):
    model = Student
    fields = ['name', 'id_card_number', 'ie_program', 'date_of_document']
    # template_name = 'icp/student_update.html'

    def get_success_url(self):
        return reverse_lazy('student_detail', kwargs={'pk': self.object.id})

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Update Student'
        return context

    def get_form_class(self):
        return modelform_factory(
            self.model,
            fields=self.fields,
            widgets={
                "date_of_document": DatePickerInput(
                    options=FlatpickrOptions(minDate="today"),
                ),
            },
        )

class StudentDeleteView(ODDeleteView):
    model = Student
    success_url = reverse_lazy('list_icps')

class GeneralInformationView(ODCreateView):
    model = GeneralInformation
    fields = ['parent_concerns', 'medical_alerts']

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'General Information'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('learning_profile', kwargs={'student_id': self.kwargs['student_id']})

class LearningProfileView(ODCreateView):
    model = LearningProfile
    form_class = LearningProfileForm

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Learning Profile'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('developmental_areas', kwargs={'student_id': self.kwargs['student_id']})

class DevelopmentalAreasView(ODCreateView):
    model = DevelopmentalArea
    fields = [
        'cognitive', 'social_emotional', 'physical', 'language',
        'speaking', 'reading', 'writing', 'listening', 'comprehension', 'others'
    ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Developmental Areas'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('skills_strengths', kwargs={'student_id': self.kwargs['student_id']})

class SkillsStrengthsView(ODCreateView):
    model = SkillsStrengths
    fields = [
        'english', 'dhivehi', 'math', 'quran_arabic', 'islam',
        'key_competencies', 'social_skills', 'adaptive_behavior',
        'language_communication', 'others'
    ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Skills & Strengths'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('accessible_learning_support', kwargs={'student_id': self.kwargs['student_id']})

class AccessibleLearningSupportView(ODCreateView):
    model = AccessibleLearningSupport
    fields = [
        'learning_environment', 'learning_materials', 'adult_support',
        'content', 'methodology', 'assessment_modification'
    ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Accessible Learning Support'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('measurable_goals', kwargs={'student_id': self.kwargs['student_id']})

class MeasurableGoalsView(ODCreateView):
    model = Goal
    fields = [
        'area', 'annual_goal', 'assessment_criteria',
        'procedures', 'assessment_schedule', 'responsible_person'
    ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Measurable Goals'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('goal_added', kwargs={'student_id': self.kwargs['student_id']})

def goal_added(request, student_id):
    return render(request, 'icp/goal_added.html', {'student_id': student_id})

class InterventionServicesView(ODCreateView):
    model = InterventionService
    fields = [
        'program', 'frequency', 'duration', 'location', 'initiation_date'
    ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Intervention Services'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('supplementary_services', kwargs={'student_id': self.kwargs['student_id']})

class SupplementaryServicesView(ODCreateView):
    model = SupplementaryService
    fields = ['type_of_support', 'time', 'duration']

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Supplementary Services'
        return context

    def form_valid(self, form):
        form.instance.student = get_object_or_404(Student, id=self.kwargs['student_id'])
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('generate_icp', kwargs={'student_id': self.kwargs['student_id']})

def generate_icp(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    document = Document()

    # Add content to the document based on the student's ICP data
    document.add_heading(f'Individualized Care Plan for {student.name}', 0)

    # General Information
    document.add_heading('General Information', level=1)
    document.add_paragraph(f'ID Card Number: {student.id_card_number}')
    document.add_paragraph(f'IE Program: {student.ie_program}')
    
    try:
        general_info = student.general_information
        document.add_paragraph(f'Parent Concerns: {general_info.parent_concerns}')
        document.add_paragraph(f'Medical Alerts: {general_info.medical_alerts}')
    except Student.general_information.RelatedObjectDoesNotExist:
        document.add_paragraph('General Information not available')

    # Learning Profile
    document.add_heading('Learning Profile', level=1)
    learning_profile = student.learning_profile
    for field in learning_profile._meta.get_fields():
        if field.name != 'id' and field.name != 'student':
            value = getattr(learning_profile, field.name)
            if isinstance(value, bool):
                document.add_paragraph(f'{field.verbose_name}: {"Yes" if value else "No"}')
            else:
                document.add_paragraph(f'{field.verbose_name}: {value}')

    # Add other sections (Developmental Areas, Skills & Strengths, etc.)

    # Save the document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={student.name}_ICP.docx'
    document.save(response)

    return response

def icp_success(request):
    return render(request, 'icp/success.html')


